#!/usr/bin/env python3
"""Generate docs/Grokbit-Features-and-Use-Cases.docx from shipped product facts.

Sources (read by hand when authoring this script's content):
  README.md, package.json, docs/grokbit-workflows.md, docs/architecture.md
Do not invent capabilities — content mirrors those sources.
"""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Grokbit-Features-and-Use-Cases.docx"
VERSION = "2026.8.27"  # package.json display version at generation time
VERIFIED = date.today().isoformat()


def set_run_font(run, *, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size={1: 18, 2: 14, 3: 12}.get(level, 11))
    return h


def add_para(doc, text, *, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, size=size)
    return p


def add_bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r0 = p.add_run(bold_lead)
        set_run_font(r0, bold=True, size=11)
        r1 = p.add_run(text)
        set_run_font(r1, size=11)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=10)
    doc.add_paragraph()
    return table


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Grokbit")
    set_run_font(r, bold=True, size=28, color=RGBColor(0x1A, 0x1A, 0x2E))

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Features and Use Cases")
    set_run_font(r, bold=True, size=16, color=RGBColor(0x33, 0x33, 0x55))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Product document · Extension version {VERSION} · Verified {VERIFIED}\n"
        "VS Code UI for xAI Grok Build and optional Claude Code\n"
        "Not affiliated with or endorsed by xAI · MIT License"
    )
    set_run_font(r, size=9, italic=True, color=RGBColor(0x55, 0x55, 0x55))

    add_para(
        doc,
        "This document describes what Grokbit ships today. Claims are derived from the "
        "extension README, package identity, and workflow reference — not from planned "
        "or aspirational work. Where a limit is known, it is stated explicitly.",
        italic=True,
        size=10,
    )

    # 1. Product summary
    add_heading(doc, "1. What Grokbit is", 1)
    add_para(
        doc,
        "Grokbit is a VS Code extension that provides a friendly editor UI for xAI’s "
        "Grok Build CLI and, optionally, Claude Code. Each chat runs in its own native "
        "editor tab. You plan safely, review edits with inline diffs, attach files, "
        "paste screenshots, run a built-in explore → plan → implement → test → document "
        "workflow, generate images and video (Grok subscription features), and dictate "
        "by voice — without living in the terminal.",
    )
    add_para(
        doc,
        "The extension is intentionally thin: the CLI owns sessions, memory, tools, and "
        "models; Grokbit mediates file access, terminals, permissions, and the UI. "
        "Plan first is the main client-side safety layer — workspace writes and "
        "non-read-only commands are blocked until you approve a plan.",
    )
    add_bullet(doc, "Marketplace: grokbit.grokbit (publisher Grokbit)")
    add_bullet(doc, "Repository: https://github.com/irichner/grokbit-vscode")
    add_bullet(doc, "Based on phuryn/grok-build-vscode (MIT); rebranded and extended")

    # 2. Who it's for
    add_heading(doc, "2. Who this is for", 1)
    add_bullet(
        doc,
        " prefer a GUI over the terminal — chat, approve changes, and manage history "
        "without memorizing CLI commands.",
        bold_lead="Developers who",
    )
    add_bullet(
        doc,
        " a safety net — Plan first mode, permission cards, and inline diffs so you "
        "always see what will change before it lands.",
        bold_lead="Teams that want",
    )
    add_bullet(
        doc,
        " in VS Code, Cursor, Windsurf, or VSCodium — the agent sits next to your files, "
        "not in a separate app.",
        bold_lead="People who work",
    )
    add_bullet(
        doc,
        " Grok Build and/or Claude Code — one extension, per-tab choice of agent.",
        bold_lead="Users of",
    )
    add_para(
        doc,
        "Engineers can still use the CLI directly; the extension does not remove power "
        "features — it wraps them in plain language.",
    )

    # 3. Requirements
    add_heading(doc, "3. Requirements", 1)
    add_table(
        doc,
        ["Need", "Detail"],
        [
            (
                "Editor",
                "VS Code 1.94+ or a compatible editor (Cursor, Windsurf, VSCodium).",
            ),
            (
                "Grok CLI",
                "Grok Build CLI (grok) on macOS, Linux, or native Windows (WSL optional).",
            ),
            (
                "Grok auth",
                "SuperGrok or X Premium+ (grok /login), or an xAI API key. Free Grok tier does not include the CLI agent.",
            ),
            (
                "Claude (optional)",
                "Claude Code / claude CLI signed in, plus the ACP adapter (onboarding can Install it).",
            ),
            (
                "Voice (optional)",
                "ffmpeg on PATH and a separate xAI API key for Speech-to-Text (pay-as-you-go; not covered by SuperGrok alone).",
            ),
        ],
    )

    # 4. Feature catalog
    add_heading(doc, "4. Feature catalog", 1)

    add_heading(doc, "4.1 Sessions and layout", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Native session tabs",
                "Each chat is its own editor tab; several open at once; full-width canvas by default.",
            ),
            (
                "Activity-bar launcher",
                "Recent sessions (default last 30 days), status dots, rename/delete, split New for Grok or Claude.",
            ),
            (
                "Session setup",
                "Top-bar chip (and composer model chip): Agent, Model, Thinking depth, Mode. Empty-tab changes are free.",
            ),
            (
                "Resumable history",
                "Resume, rename, delete, search, paginated load (100 at a time), clear-all with open tabs protected. Merged Grok + Claude list with backend badges.",
            ),
            (
                "Tab restore",
                "After a VS Code reload, tabs restore; the visible tab reconnects first.",
            ),
            (
                "Empty-session cleanup",
                "Abandoned primer-only New session tabs do not clutter history (Grok).",
            ),
        ],
    )

    add_heading(doc, "4.2 Two agents", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Grok Build (default)",
                "xAI coding agent via the Grok CLI — models, thinking depth, plan primer, /imagine, and more.",
            ),
            (
                "Claude Code (optional)",
                "Per-tab backend switch; New Claude session; own auth/onboarding; history resumes on the correct backend.",
            ),
            (
                "Per-backend logout",
                "Log out of Grok or Claude without signing out of the other.",
            ),
        ],
    )

    add_heading(doc, "4.3 Safety and control", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Agent mode",
                "Help right away; may ask before edits or commands.",
            ),
            (
                "Plan first",
                "Draft a plan before anything changes; client-side gate blocks workspace writes and non-read-only commands until you approve.",
            ),
            (
                "Auto accept",
                "Apply edits without per-change prompts when you trust the session.",
            ),
            (
                "Permission cards + inline diffs",
                "Green/red diff inside the chat card — Allow this change / Allow always / Don’t allow. No focus-stealing diff editor tab.",
            ),
            (
                "Path/command-bound grants",
                "Allow always stays scoped to the path or command when the payload carries one.",
            ),
            (
                "Inline questions",
                "Multiple-choice cards from the agent; collapse to a green check after answer; replay on resume.",
            ),
            (
                "Remembered mode",
                "Last non-plan mode (Agent / Auto accept) remembered for new sessions; Plan first is always deliberate.",
            ),
        ],
    )

    add_heading(doc, "4.4 Grokbit Actions (bundled workflow)", 2)
    add_para(
        doc,
        "Ships with the extension and works on both Grok and Claude. Click a tile to seed "
        "the composer — nothing sends until you press Enter (or your send key).",
    )
    add_table(
        doc,
        ["Skill", "Slash", "Purpose"],
        [
            (
                "Explore",
                "/grokbit-explore",
                "Map relevant code first — compact, cited orientation before changes.",
            ),
            (
                "Plan",
                "/grokbit-plan",
                "Grounded, reviewed plan — claims from the repo, tasks with proof commands.",
            ),
            (
                "Implement",
                "/grokbit-implement",
                "One task at a time; each check passes or is reverted — no half-finished state.",
            ),
            (
                "Test",
                "/grokbit-test",
                "Baseline behavior before the change, then prove what did and didn’t change.",
            ),
            (
                "Document",
                "/grokbit-document",
                "Write the doc, then verify commands and paths in it.",
            ),
            (
                "Ship (orchestrator)",
                "/grokbit-ship",
                "Full pipeline with a mandatory human checkpoint after plan.",
            ),
        ],
    )
    add_para(
        doc,
        "User Workflows also appear when present: Grok uses .grok/workflows/*.rhai; "
        "Claude uses .claude/workflows/*.js (formats are not interchangeable). On Grok, "
        "Create Workflow opens a form + phase canvas that seeds /create-workflow. "
        "Default Actions scope is workflow-only; set grok.actionsScope to all to also "
        "list discovered skills, agents, and CLI commands. Slash “/” autocomplete always "
        "includes CLI commands. Skills provision into ~/.grok/skills and ~/.claude/skills "
        "on activation (grok.skills.provision).",
    )

    add_heading(doc, "4.5 Context and files", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "File chips",
                "Drag from Explorer, + button, right-click Grokbit: Send File, or Alt+G.",
            ),
            (
                "Active-editor context",
                "Optional auto-include of the open file (grok.includeActiveFileByDefault). Explicit attachments rank stronger.",
            ),
            (
                "Shift-drag embed",
                "Embed file contents as a fenced block instead of a path reference.",
            ),
            (
                "Paste screenshots",
                "Paste into the composer (thumbnail above input). Vision-capable agents may receive the image; Grok Build currently gets a path plus a notice when it cannot view images.",
            ),
            (
                "Docs browser",
                "Top-bar Docs — browse business documents in the workspace; open, reveal, attach, or seed into the composer.",
            ),
            (
                "Document cards",
                "When a turn produces a doc path — Copy / Open / Reveal. Generation via CLI skills (/docx, /pptx, /xlsx, …) or plain language — not a built-in Office suite.",
            ),
        ],
    )

    add_heading(doc, "4.6 While the agent works", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Activity carousel",
                "One compact strip per turn (current action, step counter, peek); collapses to a plain-language summary. Off = classic stream (grok.compactActivity).",
            ),
            (
                "Tool rows",
                "Category icons, expandable command output + copy, failed tools with a short reason, view-diff on past edits.",
            ),
            (
                "Files-changed strip",
                "Scannable auth.ts +12 −3 chips above the composer for this turn’s applied edits.",
            ),
            (
                "Thinking control",
                "Full traces off by default; Thinking… stand-in; toggle live (grok.showThinking).",
            ),
            (
                "Streaming answers",
                "Agent text streams live with markdown, code, and tables.",
            ),
        ],
    )

    add_heading(doc, "4.7 Awareness and multi-session", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Status dots",
                "Blue working · yellow needs you · green/red unread done · gray at rest — on launcher and history rows.",
            ),
            (
                "Status-bar HUD",
                "Always-visible model · thinking · mode · context % · working spinner · needs-you count. Click to jump to chat.",
            ),
            (
                "Background notifications",
                "Optional passive Go to session when a hidden tab needs you (grok.notifyWhenWaiting — never steals focus).",
            ),
        ],
    )

    add_heading(doc, "4.8 Models, media, and voice", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            (
                "Model picker",
                "Choose model from setup chip, model chip, gear, or command palette. Cross-agent switches restart cleanly.",
            ),
            (
                "Thinking depth",
                "Per-tab reasoning effort (none → xhigh) for Grok; omitted cleanly for Claude (no effort axis).",
            ),
            (
                "Image and video",
                "/imagine and /imagine-video render inline (subscription Grok features); copy path / open; resume-safe. Reference-photo edit when the CLI offers it.",
            ),
            (
                "Voice control",
                "Mic in the composer → live STT; say “grok send” (configurable) to submit. Needs ffmpeg + separate xAI API key.",
            ),
        ],
    )

    add_heading(doc, "4.9 Rich chat and power tools", 2)
    add_table(
        doc,
        ["Capability", "What you get"],
        [
            ("Markdown", "Full formatting — code, lists, tables."),
            (
                "LaTeX",
                "MathJax offline; hover to copy or export PNG/SVG.",
            ),
            (
                "Mermaid",
                "Inline diagrams (theme-aware); hover to copy or export.",
            ),
            (
                "Slash commands",
                "Type / for CLI autocomplete (/compact, /imagine, skills, …).",
            ),
            (
                "MCP tools",
                "Configured in the CLI config; gear links to config; new session to reload.",
            ),
            (
                "Context donut",
                "How full the conversation window is; /compact or new session to free room.",
            ),
            (
                "Chat zoom",
                "Zoom only the chat panel 60–300% (grok.chatFontScale).",
            ),
            (
                "Keyboard shortcuts",
                "Ctrl/Cmd+; open, Alt+G attach file, Enter/Ctrl+Enter send — listed in gear → Keyboard shortcuts.",
            ),
            (
                "Telemetry",
                "Anonymous session_start only (install id + mode/model/effort) — never content, code, or paths. Opt out anytime.",
            ),
        ],
    )

    # 5. Use cases
    add_heading(doc, "5. Use cases", 1)
    add_para(
        doc,
        "Each use case maps a real job to shipped capabilities. Steps describe what a "
        "user does in the product, not internal architecture.",
    )

    add_heading(doc, "UC-1 — First coding session in the editor", 2)
    add_para(doc, "Goal: Get help changing code without leaving VS Code.", bold=True)
    add_bullet(doc, "Install the Grok CLI, sign in (subscription or API key), install the Grokbit extension.")
    add_bullet(doc, "Open Grokbit (activity bar or Ctrl/Cmd+;), start New session.")
    add_bullet(doc, "Optionally set Agent / Model / Thinking / Mode on the Session setup chip.")
    add_bullet(doc, "Describe the task in plain language; attach files with chips or Alt+G.")
    add_bullet(doc, "Review permission cards and inline diffs before edits land.")
    add_para(
        doc,
        "Outcome: A streaming answer, tool activity in the carousel, and applied edits "
        "summarized in the files-changed strip — with an audit trail you can resume later.",
        italic=True,
    )

    add_heading(doc, "UC-2 — Plan before any file changes", 2)
    add_para(doc, "Goal: Agree on an approach before the agent touches the repo.", bold=True)
    add_bullet(doc, "Switch Mode to Plan first (or open a session already in Plan).")
    add_bullet(doc, "Describe the change; optionally run /grokbit-plan for a structured plan artifact.")
    add_bullet(doc, "Read the plan card; Approve, Keep planning, or Cancel.")
    add_bullet(
        doc,
        "Until you approve, the client gate blocks workspace writes and non-read-only commands.",
    )
    add_para(
        doc,
        "Outcome: No workspace mutation until an explicit plan verdict; implementation can follow in Agent or Auto accept.",
        italic=True,
    )

    add_heading(doc, "UC-3 — Full feature with the bundled pipeline", 2)
    add_para(
        doc,
        "Goal: Explore → plan → implement → test → document with human control after planning.",
        bold=True,
    )
    add_bullet(doc, "From Grokbit Actions, click Explore (optional) then Plan — or invoke /grokbit-ship.")
    add_bullet(doc, "Review the plan at the mandatory checkpoint; refine or approve.")
    add_bullet(doc, "Run Implement (task-by-task verify-or-revert) and Test (baseline + verify).")
    add_bullet(doc, "Run Document when the change needs a durable guide or reference.")
    add_para(
        doc,
        "Outcome: A disciplined change path that works the same on Grok and Claude tabs; "
        "artifacts under .grokbit/plans/ when Plan is used.",
        italic=True,
    )

    add_heading(doc, "UC-4 — Parallel work across sessions", 2)
    add_para(doc, "Goal: Keep several agent conversations open without losing state.", bold=True)
    add_bullet(doc, "Open multiple session tabs; use the launcher Recent list and status dots.")
    add_bullet(doc, "Watch the status-bar HUD for model, mode, context %, and needs-you count.")
    add_bullet(doc, "Optionally enable grok.notifyWhenWaiting for passive alerts on background tabs.")
    add_bullet(doc, "Resume any past session from history; Grok and Claude rows resume on the correct backend.")
    add_para(
        doc,
        "Outcome: Multi-threaded agent work with clear signals when a tab needs your approval.",
        italic=True,
    )

    add_heading(doc, "UC-5 — Switch between Grok and Claude for the same project", 2)
    add_para(doc, "Goal: Use either agent without installing a second chat UI.", bold=True)
    add_bullet(doc, "Use the launcher split New button or Session setup Agent control.")
    add_bullet(doc, "Complete Claude onboarding (adapter Install + auth) when first needed.")
    add_bullet(doc, "History lists both backends with badges; resume preserves backend.")
    add_bullet(doc, "Log out of one backend without signing out of the other.")
    add_para(
        doc,
        "Outcome: One extension, per-tab agent choice, shared project history list.",
        italic=True,
    )

    add_heading(doc, "UC-6 — Generate media or dictate a prompt", 2)
    add_para(doc, "Goal: Produce image/video assets or speak a prompt hands-free.", bold=True)
    add_bullet(doc, "Image/video: use /imagine or /imagine-video (Grok subscription); results render inline.")
    add_bullet(doc, "Voice: configure API key + ffmpeg; use the mic; say the send phrase to submit.")
    add_para(
        doc,
        "Outcome: Inline media with open/copy actions; voice as an alternate input path "
        "(STT is pay-as-you-go and separate from SuperGrok).",
        italic=True,
    )

    add_heading(doc, "UC-7 — Review what the agent did this turn", 2)
    add_para(doc, "Goal: Understand tool use and edits without scrolling a noisy log.", bold=True)
    add_bullet(doc, "Leave Compact activity on (default) for a single carousel strip per turn.")
    add_bullet(doc, "Expand for tool rows, command output, and view-diff on past edits.")
    add_bullet(doc, "Use the files-changed strip above the composer to jump to applied edits.")
    add_para(
        doc,
        "Outcome: Plain-language summaries (e.g. Explored N items, edited M files) with drill-down detail.",
        italic=True,
    )

    add_heading(doc, "UC-8 — Attach the right context for a hard bug", 2)
    add_para(doc, "Goal: Point the agent at specific files and screenshots.", bold=True)
    add_bullet(doc, "Attach via chips, drag-drop, Send File, or Alt+G; Shift-drag to embed contents.")
    add_bullet(doc, "Paste a screenshot into the composer when the agent can use vision.")
    add_bullet(doc, "Use Docs to attach an existing workspace document path.")
    add_para(
        doc,
        "Outcome: Explicit attachments outrank ambient open-file context; the prompt envelope "
        "keeps user intent clear.",
        italic=True,
    )

    # 6. Modes quick reference
    add_heading(doc, "6. Modes at a glance", 1)
    add_table(
        doc,
        ["Mode", "In plain English"],
        [
            (
                "Agent (default)",
                "The agent can help right away. It may ask before editing files or running commands.",
            ),
            (
                "Plan first",
                "The agent drafts a plan first. Nothing changes in your project until you approve the plan.",
            ),
            (
                "Auto accept",
                "The agent acts without asking for permission on each change. Use when you trust the session.",
            ),
        ],
    )

    # 7. Known limits
    add_heading(doc, "7. Known limits (honest)", 1)
    add_bullet(
        doc,
        " compares proposed old vs new text (not necessarily disk at preview time); the write happens after you approve.",
        bold_lead="Diff preview",
    )
    add_bullet(doc, " UI is not shipped yet.", bold_lead="Worktree")
    add_bullet(
        doc,
        " inspector is not shipped (current CLI builds do not expose nested subagent cards over ACP).",
        bold_lead="Subagent",
    )
    add_bullet(
        doc,
        " defaults to the bundled suite plus User Workflows; set grok.actionsScope to all for skills/agents/commands in that menu.",
        bold_lead="Grokbit Actions",
    )
    add_bullet(
        doc,
        " User Workflows are backend-native (Rhai vs JS) — not interchangeable across Grok and Claude.",
        bold_lead="Formats:",
    )
    add_bullet(
        doc,
        " defaults to the left activity bar; move to the secondary side bar manually if you prefer the right.",
        bold_lead="Launcher",
    )
    add_bullet(
        doc,
        " mid-turn auto-cards for every business-doc path are disabled (flooding coding sessions); Docs browser and path classification remain.",
        bold_lead="Business document",
    )

    # 8. Privacy
    add_heading(doc, "8. Privacy", 1)
    add_para(
        doc,
        "Telemetry is privacy-by-design: no message content, no code, no file paths, and no "
        "account identity leave your machine via telemetry. One anonymous session_start per "
        "real conversation carries install id + mode/model/effort only. Turn off with "
        "grok.telemetry.enabled: false or VS Code’s global telemetry setting. Details: docs/privacy.md.",
    )

    # 9. Further reading
    add_heading(doc, "9. Further reading in the repository", 1)
    add_table(
        doc,
        ["Document", "Path"],
        [
            ("Product README (canonical feature prose)", "README.md"),
            ("Architecture map", "docs/architecture.md"),
            ("Grokbit workflow deep dive", "docs/grokbit-workflows.md"),
            ("Slash command snapshot", "docs/SLASH-COMMANDS.md"),
            ("Privacy", "docs/privacy.md"),
            ("Test plan / taxonomy", "TESTS.md"),
            ("Project map for contributors", "CLAUDE.md"),
        ],
    )

    add_heading(doc, "10. Provenance", 1)
    add_para(
        doc,
        "Derived from README.md (capabilities overview, who this is for, requirements, "
        "features in depth, known limits, privacy), package.json (name, version, description, "
        "publisher), and docs/grokbit-workflows.md (pipeline skills). Generated as a Word "
        f"export for offline sharing. Verified date: {VERIFIED}. Extension version stamp: {VERSION}.",
        size=10,
    )
    add_para(
        doc,
        "Note: docs/FEATURES.md and docs/USER_GUIDE.md in this repository currently describe "
        "a Claude Code agentic template, not the Grokbit extension product. This document "
        "intentionally does not use them as sources.",
        size=10,
        italic=True,
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        "Grokbit · MIT · https://marketplace.visualstudio.com/items?itemName=grokbit.grokbit"
    )
    set_run_font(r, size=9, color=RGBColor(0x66, 0x66, 0x66))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    path = build()
    print(path)
